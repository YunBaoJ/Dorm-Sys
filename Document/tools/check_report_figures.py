from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "开题报告-孙乾云.docx"

doc = Document(DOCX)
print(f"inline_shapes={len(doc.inline_shapes)}")
for i, shape in enumerate(doc.inline_shapes, start=1):
    print(f"{i}: {shape.width.cm:.2f}cm x {shape.height.cm:.2f}cm")

cell = doc.tables[0].cell(3, 1)
for i, para in enumerate(cell.paragraphs):
    text = para.text.strip()
    has_pic = bool(para._p.xpath(".//w:drawing"))
    if has_pic or text.startswith("图1") or text.startswith("图2"):
        print(f"{i}: pic={has_pic} align={para.alignment} text={text}")

print("table col widths")
for i, col in enumerate(doc.tables[0].columns):
    print(i, [cell.width.cm if cell.width else None for cell in col.cells[:5]])
