from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "开题报告-孙乾云.docx"

doc = Document(DOCX)
keys = ("图1", "图2", "图一", "图二", "开发与部署流程", "系统架构图", "技术路线")

print("PARAGRAPHS")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if any(key in text for key in keys):
        print(f"{i}: {text}")

print("TABLES")
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            if text and any(key in text for key in keys):
                print(f"table {ti} row {ri} col {ci}: {text}")

print("CELL PARAGRAPHS AROUND FIGURES")
cell = doc.tables[0].cell(3, 1)
for i, para in enumerate(cell.paragraphs):
    text = para.text.strip()
    if text:
        marker = "*" if any(key in text for key in keys) else " "
        safe = text.encode("unicode_escape").decode("ascii")
        print(f"{marker}{i}: {safe[:260]}")
