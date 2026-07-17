from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "开题报告-孙乾云.docx"
IMG1 = ROOT / "images" / "图1-开发与部署流程.png"
IMG2 = ROOT / "images" / "图2-系统架构图.png"


def insert_paragraph_before(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def clear_paragraph(paragraph):
    for run in paragraph.runs:
        run.text = ""
    paragraph._p.clear_content()


def remove_existing_picture_paragraphs(cell):
    for paragraph in list(cell.paragraphs):
        has_pic = bool(paragraph._p.xpath(".//pic:pic"))
        has_drawing = bool(paragraph._p.xpath(".//w:drawing"))
        if has_pic or has_drawing:
            paragraph._element.getparent().remove(paragraph._element)


def add_centered_picture_before(paragraph, image_path, width_cm):
    p = insert_paragraph_before(paragraph)
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(0.15)
    p.paragraph_format.space_after = Cm(0.12)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return p


def center_caption(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Cm(0.05)
    paragraph.paragraph_format.space_after = Cm(0.12)
    for run in paragraph.runs:
        run.bold = False


doc = Document(DOCX)
cell = doc.tables[0].cell(3, 1)
remove_existing_picture_paragraphs(cell)

targets = {
    "图1   开发与部署流程": (IMG1, 12.8),
    "图2   系统架构图": (IMG2, 13.2),
}

inserted = set()
for paragraph in list(cell.paragraphs):
    text = paragraph.text.strip()
    if text in targets and text not in inserted:
        image_path, width_cm = targets[text]
        add_centered_picture_before(paragraph, image_path, width_cm)
        center_caption(paragraph)
        inserted.add(text)

missing = set(targets) - inserted
if missing:
    raise SystemExit(f"未找到图题：{', '.join(sorted(missing))}")

doc.save(DOCX)
print(DOCX)
